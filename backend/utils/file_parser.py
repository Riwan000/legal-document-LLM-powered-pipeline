"""
File parsing utilities for PDF and DOCX documents.
Extracts text with page number tracking.
Supports OCR for scanned/image-based PDFs.
Uses batched OCR and configurable DPI to avoid MemoryError in subprocess reader.
"""
import os
from pathlib import Path
from typing import List, Tuple, Optional, Union
import pypdf
from backend.models.document_structure import PageInfo
from docx import Document
from backend.config import settings

# Optional OCR imports (for scanned PDFs)
try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


class FileParser:
    """Parser for PDF and DOCX files with page tracking."""
    
    @staticmethod
    def parse_pdf(file_path: Path) -> List[Tuple[str, int]]:
        """Parse PDF and return (text, page_number) tuples. Use parse_pdf_pages() for PageInfo with is_ocr."""
        pages = FileParser.parse_pdf_pages(file_path)
        return [(p.text, p.page_number) for p in pages]

    @staticmethod
    def parse_pdf_pages(file_path: Path) -> List[PageInfo]:
        """
        Parse PDF file and extract text with page numbers and OCR flag.
        Falls back to OCR if text extraction fails (for scanned PDFs).
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            List of PageInfo (text, page_number, is_ocr)
        """
        text_pages: List[PageInfo] = []
        pages_needing_ocr: List[int] = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                for page_num in range(total_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    text = text.strip() if text else ""
                    
                    if text:
                        text_pages.append(PageInfo(text=text, page_number=page_num + 1, is_ocr=False))
                    else:
                        pages_needing_ocr.append(page_num + 1)
                        
        except Exception as e:
            raise ValueError(f"Error parsing PDF {file_path}: {str(e)}")
        
        if pages_needing_ocr and OCR_AVAILABLE:
            try:
                ocr_results = FileParser._parse_pdf_with_ocr_selective(file_path, pages_needing_ocr)
                text_pages.extend(ocr_results)
            except Exception:
                if not text_pages:
                    try:
                        ocr_pages = FileParser._parse_pdf_with_ocr(file_path)
                        text_pages.extend(ocr_pages)
                    except Exception as full_ocr_error:
                        print(f"Warning: OCR extraction failed: {str(full_ocr_error)}")
                        print("Note: Install Tesseract OCR for scanned PDF support.")
        
        if not text_pages:
            print(f"Warning: No text extracted from PDF {file_path}")
        
        return text_pages
    
    @staticmethod
    def _preprocess_for_ocr(image):
        """Grayscale → contrast boost → sharpen → binarize before Tesseract."""
        from PIL import ImageFilter, ImageEnhance
        image = image.convert("L")
        image = ImageEnhance.Contrast(image).enhance(2.0)
        image = image.filter(ImageFilter.SHARPEN)
        image = image.point(lambda x: 0 if x < 140 else 255, '1')
        return image

    @staticmethod
    def _ocr_confidence(image, lang: str) -> float:
        """Return average word-level confidence (0-100) from Tesseract."""
        data = pytesseract.image_to_data(image, lang=lang, output_type=pytesseract.Output.DICT)
        confs = [int(c) for c in data['conf'] if str(c) != '-1' and int(c) >= 0]
        return round(sum(confs) / len(confs), 1) if confs else -1.0

    @staticmethod
    def _parse_pdf_with_ocr_selective(file_path: Path, page_numbers: List[int]) -> List[PageInfo]:
        """
        Parse specific pages of PDF using OCR (for pages that failed text extraction).

        Args:
            file_path: Path to PDF file
            page_numbers: List of page numbers (1-indexed) to process with OCR

        Returns:
            List of PageInfo with is_ocr=True and ocr_confidence populated.
        """
        if not OCR_AVAILABLE:
            return []

        text_pages: List[PageInfo] = []

        try:
            dpi = getattr(settings, "OCR_DPI", 200)
            images = convert_from_path(str(file_path), dpi=dpi, first_page=min(page_numbers), last_page=max(page_numbers))
            page_map = {i: page_num for i, page_num in enumerate(range(min(page_numbers), max(page_numbers) + 1), start=0)}

            # Extract text from specified pages using OCR (one page at a time to limit subprocess buffer)
            for img_idx, image in enumerate(images):
                actual_page = page_map.get(img_idx)
                if actual_page and actual_page in page_numbers:
                    try:
                        conf = FileParser._ocr_confidence(image, lang=settings.OCR_LANGUAGE)
                        preprocessed = FileParser._preprocess_for_ocr(image)
                        custom_config = r'--oem 3 --psm 6'
                        text = pytesseract.image_to_string(preprocessed, lang=settings.OCR_LANGUAGE, config=custom_config)
                        text = text.strip()
                        if text:
                            text_pages.append(PageInfo(text=text, page_number=actual_page, is_ocr=True, ocr_confidence=conf))
                    except Exception as e:
                        print(f"Warning: OCR failed for page {actual_page}: {str(e)}")
                del image  # release image before next to reduce peak memory

        except Exception as e:
            print(f"Warning: Selective OCR failed: {str(e)}")

        return text_pages
    
    @staticmethod
    def _parse_pdf_with_ocr(file_path: Path) -> List[PageInfo]:
        """
        Parse PDF using OCR (for scanned/image-based PDFs).
        Processes in batches and uses configurable DPI to avoid MemoryError
        in subprocess reader (buffer.append(fh.read())).
        """
        if not OCR_AVAILABLE:
            raise ImportError("OCR libraries not available. Install pytesseract, pdf2image, and Pillow.")

        text_pages: List[PageInfo] = []
        dpi = getattr(settings, "OCR_DPI", 200)
        batch_size = getattr(settings, "OCR_PAGES_PER_BATCH", 10)

        try:
            with open(file_path, "rb") as f:
                total_pages = len(pypdf.PdfReader(f).pages)
        except Exception as e:
            raise ValueError(f"Error reading PDF {file_path}: {e}") from e

        first_page = 1
        while first_page <= total_pages:
            last_page = min(first_page + batch_size - 1, total_pages)
            try:
                images = convert_from_path(
                    str(file_path),
                    dpi=dpi,
                    first_page=first_page,
                    last_page=last_page,
                )
            except FileNotFoundError as e:
                if "poppler" in str(e).lower() or "pdftoppm" in str(e).lower():
                    raise RuntimeError(
                        "Poppler not found. Required for PDF to image conversion:\n"
                        "Windows: Download from https://github.com/oschwartz10612/poppler-windows/releases\n"
                        "Linux: sudo apt-get install poppler-utils\n"
                        "macOS: brew install poppler"
                    ) from e
                raise ValueError(f"Error performing OCR on PDF {file_path}: {e}") from e
            try:
                for idx, image in enumerate(images):
                    page_num = first_page + idx
                    try:
                        conf = FileParser._ocr_confidence(image, lang=settings.OCR_LANGUAGE)
                        preprocessed = FileParser._preprocess_for_ocr(image)
                        custom_config = r'--oem 3 --psm 6'
                        text = pytesseract.image_to_string(preprocessed, lang=settings.OCR_LANGUAGE, config=custom_config)
                        text = text.strip()
                        if text:
                            text_pages.append(PageInfo(text=text, page_number=page_num, is_ocr=True, ocr_confidence=conf))
                    except pytesseract.TesseractNotFoundError:
                        raise RuntimeError(
                            "Tesseract OCR not found. Please install Tesseract:\n"
                            "Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki\n"
                            "Linux: sudo apt-get install tesseract-ocr\n"
                            "macOS: brew install tesseract"
                        )
                    finally:
                        del image
            finally:
                del images
            first_page = last_page + 1

        return text_pages
    
    @staticmethod
    def parse_docx(file_path: Path) -> List[Tuple[str, int]]:
        """Parse DOCX and return (text, page_number) tuples. Use parse_docx_pages() for PageInfo."""
        pages = FileParser.parse_docx_pages(file_path)
        return [(p.text, p.page_number) for p in pages]

    @staticmethod
    def parse_docx_pages(file_path: Path) -> List[PageInfo]:
        """
        Parse DOCX file and extract text with estimated page numbers.
        DOCX has no native page numbers; is_ocr is always False.
        
        Returns:
            List of PageInfo (text, estimated_page_number, is_ocr=False)
        """
        text_pages: List[PageInfo] = []
        
        try:
            doc = Document(file_path)
            words_per_page = 500
            current_text = ""
            current_page = 1
            word_count = 0
            
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    current_text += para_text + "\n"
                    word_count += len(para_text.split())
                    if word_count >= words_per_page:
                        text_pages.append(PageInfo(text=current_text.strip(), page_number=current_page, is_ocr=False))
                        current_text = ""
                        word_count = 0
                        current_page += 1
            
            if current_text.strip():
                text_pages.append(PageInfo(text=current_text.strip(), page_number=current_page, is_ocr=False))
            
            if not text_pages:
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                if full_text:
                    text_pages.append(PageInfo(text=full_text, page_number=1, is_ocr=False))
                    
        except Exception as e:
            raise ValueError(f"Error parsing DOCX {file_path}: {str(e)}")
        
        return text_pages
    
    @staticmethod
    def parse_file(file_path: Path) -> List[Tuple[str, int]]:
        """
        Parse file based on extension (PDF or DOCX).
        Returns (text, page_number) for backward compatibility.
        """
        pages = FileParser.parse_file_to_pages(file_path)
        return [(p.text, p.page_number) for p in pages]

    @staticmethod
    def parse_file_to_pages(file_path: Path) -> List[PageInfo]:
        """
        Parse file and return list of PageInfo (text, page_number, is_ocr).
        Use this when OCR flag is needed (e.g. ingestion pipeline).
        """
        file_ext = file_path.suffix.lower()
        if file_ext == '.pdf':
            return FileParser.parse_pdf_pages(file_path)
        if file_ext in ['.docx', '.doc']:
            return FileParser.parse_docx_pages(file_path)
        raise ValueError(f"Unsupported file type: {file_ext}. Supported: .pdf, .docx")

